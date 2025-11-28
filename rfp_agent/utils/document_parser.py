"""
Document parsing utilities for extracting text and metadata from various file formats.
"""

import io
from pathlib import Path
from typing import Dict, Any, Optional

import PyPDF2
from docx import Document


class DocumentParser:
    """
    Parser for extracting text and metadata from various document formats.
    """
    
    @staticmethod
    def parse_document(file_path: str) -> Dict[str, Any]:
        """
        Parse a document and extract text and metadata.
        
        Args:
            file_path: Path to the document file
        
        Returns:
            Dictionary containing text, metadata, and file info
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        parsers = {
            '.pdf': DocumentParser._parse_pdf,
            '.docx': DocumentParser._parse_docx,
            '.doc': DocumentParser._parse_doc,
            '.txt': DocumentParser._parse_txt,
        }
        
        parser = parsers.get(extension)
        if not parser:
            raise ValueError(f"Unsupported file format: {extension}")
        
        result = parser(file_path)
        result['file_info'] = {
            'name': path.name,
            'size_bytes': path.stat().st_size,
            'extension': extension,
            'path': str(path.absolute()),
        }
        
        return result
    
    @staticmethod
    def _parse_pdf(file_path: str) -> Dict[str, Any]:
        """Parse PDF document."""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            # Extract text from all pages
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())
            
            text = '\n\n'.join(text_parts)
            
            # Extract metadata
            metadata = {}
            if reader.metadata:
                metadata = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                }
            
            return {
                'text': text,
                'metadata': metadata,
                'page_count': len(reader.pages),
            }
    
    @staticmethod
    def _parse_docx(file_path: str) -> Dict[str, Any]:
        """Parse DOCX document."""
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        text = '\n\n'.join(text_parts)
        
        # Extract metadata
        metadata = {
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
            'subject': doc.core_properties.subject or '',
            'created': str(doc.core_properties.created) if doc.core_properties.created else '',
        }
        
        return {
            'text': text,
            'metadata': metadata,
            'paragraph_count': len(doc.paragraphs),
        }
    
    @staticmethod
    def _parse_doc(file_path: str) -> Dict[str, Any]:
        """Parse DOC document (legacy format)."""
        # For .doc files, we'd need additional libraries like antiword or textract
        # For now, return a placeholder
        return {
            'text': '',
            'metadata': {},
            'error': 'Legacy .doc format requires additional conversion tools',
        }
    
    @staticmethod
    def _parse_txt(file_path: str) -> Dict[str, Any]:
        """Parse plain text document."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            text = file.read()
        
        return {
            'text': text,
            'metadata': {},
            'line_count': text.count('\n') + 1,
        }
    
    @staticmethod
    def extract_client_info(text: str) -> Dict[str, Optional[str]]:
        """
        Extract client information from RFP text using heuristics.
        
        Args:
            text: Full text of the RFP document
        
        Returns:
            Dictionary with extracted client name, RFP title, etc.
        """
        import re
        
        info = {
            'client_name': None,
            'rfp_title': None,
            'rfp_number': None,
            'deadline': None,
        }
        
        # Look for common patterns
        lines = text.split('\n')[:50]  # Check first 50 lines
        
        # Try to find RFP title (often in first few lines, all caps or title case)
        for line in lines[:10]:
            line = line.strip()
            if len(line) > 20 and len(line) < 150:
                if line.isupper() or line.istitle():
                    if not info['rfp_title']:
                        info['rfp_title'] = line
        
        # Look for RFP number
        rfp_number_pattern = r'RFP[#\s-]*(\d+[-\d]*)'
        match = re.search(rfp_number_pattern, text, re.IGNORECASE)
        if match:
            info['rfp_number'] = match.group(1)
        
        # Look for deadline
        deadline_patterns = [
            r'deadline[:\s]+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
            r'due\s+date[:\s]+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
            r'submission\s+date[:\s]+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
        ]
        for pattern in deadline_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                info['deadline'] = match.group(1)
                break
        
        return info
